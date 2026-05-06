from __future__ import annotations

from docx import Document
from docx.shared import Inches, RGBColor

from app.image_utils import calculate_fit_size_inches
from app.models import Issue, Report, ScreenshotEntry, display_language


def export_report_to_docx(report: Report, output_path: str) -> list[str]:
    document = Document()
    _add_cover_page(document, report)
    skipped = _add_screenshots(document, report)
    if not report.screenshots:
        raise ValueError("No hay capturas para exportar.")
    if len(skipped) == len(report.screenshots):
        raise ValueError("No se pudo exportar ninguna captura al documento.")
    document.save(output_path)
    return skipped


def _add_cover_page(document: Document, report: Report) -> None:
    document.add_heading("Informe QA de Traduccion", 0)

    _add_label_value(document, "Juego", report.game_name)
    _add_label_value(document, "Traductor", report.translator)
    _add_label_value(document, "Tester", report.tester)
    _add_label_value(document, "Idioma original", display_language(report.source_language))
    _add_label_value(document, "Idioma destino", display_language(report.target_language))
    _add_label_value(document, "Fecha", report.report_date)

    document.add_page_break()


def _add_label_value(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(str(value or ""))


def _add_screenshots(document: Document, report: Report) -> list[str]:
    skipped: list[str] = []
    total = len(report.screenshots)

    for idx, shot in enumerate(report.screenshots):
        compact = _is_compact_entry(shot)
        is_last_capture = idx == (total - 1)

        if not _add_screenshot_block(
            document,
            shot,
            compact=compact,
            is_last_capture=is_last_capture,
        ):
            skipped.append(shot.image_path)
    return skipped


def _is_compact_entry(entry: ScreenshotEntry) -> bool:
    if len(entry.issues) > 2:
        return False

    total_chars = sum(
        len(issue.wrong_text) + len(issue.correction) + len(issue.note) for issue in entry.issues
    )
    return total_chars <= 450


def _add_screenshot_block(
    document: Document,
    entry: ScreenshotEntry,
    *,
    compact: bool,
    is_last_capture: bool,
) -> bool:
    max_width_in = 6.0
    max_height_in = 3.1 if compact else 5.0
    try:
        width_in, height_in = calculate_fit_size_inches(entry.image_path, max_width_in, max_height_in)
    except OSError:
        return False

    try:
        pic_paragraph = document.add_paragraph()
        pic_paragraph.paragraph_format.keep_with_next = True
        pic_paragraph.add_run().add_picture(
            entry.image_path, width=Inches(width_in), height=Inches(height_in)
        )
    except Exception:  # noqa: BLE001
        return False

    if not entry.issues:
        empty_paragraph = document.add_paragraph("Sin errores registrados.")
        empty_paragraph.runs[0].italic = True
        return True

    total_issues = len(entry.issues)
    for issue_idx, issue in enumerate(entry.issues, start=1):
        _add_issue_block(document, issue_idx, issue)
        if issue_idx < total_issues:
            document.add_paragraph("")

    if total_issues > 1 and not is_last_capture:
        document.add_paragraph("")

    return True


def _add_issue_block(document: Document, _issue_idx: int, issue: Issue) -> None:
    wrong_p = document.add_paragraph()
    wrong_run = wrong_p.add_run(issue.wrong_text or "(vacio)")
    wrong_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    correction_p = document.add_paragraph()
    correction_run = correction_p.add_run(issue.correction or "(vacio)")
    correction_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    if issue.note.strip():
        note_p = document.add_paragraph()
        note_run = note_p.add_run(issue.note.strip())
        note_run.font.color.rgb = RGBColor(0x00, 0x64, 0x00)
